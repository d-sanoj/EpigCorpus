#!/usr/bin/env python3
"""paper/EpigCorpus_paper_v1.md -> .docx with figures embedded and captions."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent.parent
MD = REPO/"paper"/"EpigCorpus_paper_v1.md"
FIG = REPO/"figures"
OUT = REPO/"paper"/"EpigCorpus_paper_v1.docx"

# figures inserted after the section they belong to
PLACE = {
    "4. What the plain text encodes": ["F4_vinculum_census", "F3_symbol_inventory"],
    "5. Dataset construction":                         ["F1_exclusion_surface", "F10_testset_ceilings"],
    "6. Artifacts and bias":                           ["F2_ordering_and_bias", "F5_circularity", "F9_editorial_noise"],
    "8. Results":                                      ["F6_context_conditions", "F7_memorisation_vs_transfer", "F8_cost_performance"],
}

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

def add_fig(name):
    png = FIG/f"{name}.png"; cap = FIG/f"{name}.caption.txt"
    if not png.exists(): return
    doc.add_picture(str(png), width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.exists():
        p = doc.add_paragraph(cap.read_text(encoding="utf-8").strip())
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.size = Pt(8.5); r.italic = True
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def emit_table(rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
    if not cells: return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row[:len(cells[0])]):
            val = re.sub(r"\*\*(.+?)\*\*", r"\1", val).replace("`", "")
            cell = t.cell(i, j); cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
                    if i == 0: r.bold = True
    doc.add_paragraph()

_md = MD.read_text(encoding="utf-8")
_heads = {l[3:].strip() for l in _md.split("\n") if l.startswith("## ")}
_bad = [k for k in PLACE if k not in _heads]
if _bad:
    raise SystemExit(f"figure placement keys match no heading: {_bad}\navailable: {sorted(_heads)}")
lines = _md.split("\n")
i = 0; pending = None; buf = []
while i < len(lines):
    L = lines[i]
    if L.startswith("|"):
        buf.append(L); i += 1; continue
    if buf:
        emit_table(buf); buf = []
    if L.startswith("# "):
        doc.add_heading(L[2:].strip(), 0)
    elif L.startswith("## "):
        if pending:
            for f in PLACE.get(pending, []): add_fig(f)
        pending = L[3:].strip()
        doc.add_heading(pending, 1)
    elif L.startswith("### "):
        doc.add_heading(L[4:].strip(), 2)
    elif L.strip() == "---":
        pass
    elif L.startswith("> "):
        p = doc.add_paragraph(re.sub(r"[*`]", "", L[2:]))
        p.paragraph_format.left_indent = Inches(0.35)
        for r in p.runs: r.italic = True
    elif L.strip():
        txt = L.rstrip()
        para = doc.add_paragraph()
        for part in re.split(r"(\*\*.+?\*\*)", txt):
            if part.startswith("**") and part.endswith("**"):
                para.add_run(part[2:-2].replace("`", "")).bold = True
            else:
                para.add_run(re.sub(r"[*`]", "", part))
    i += 1
if buf: emit_table(buf)
if pending:
    for f in PLACE.get(pending, []): add_fig(f)

doc.save(OUT)
print(f"wrote {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
print(f"  figures embedded: {sum(len(v) for v in PLACE.values())}")
